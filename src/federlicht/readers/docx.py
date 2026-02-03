from __future__ import annotations

from pathlib import Path

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    docx = None


def _docx_available() -> bool:
    return docx is not None


def read_docx_text(docx_path: Path, max_chars: int, start: int = 0) -> str:
    if not _docx_available():
        return "python-docx is not installed. Cannot read DOCX."
    document = docx.Document(docx_path)  # type: ignore[attr-defined]
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for idx, table in enumerate(document.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append(f"[table {idx + 1}]")
            parts.extend(rows)
    text = "\n".join(parts)
    if start > 0:
        text = text[start:]
    if max_chars and max_chars > 0 and len(text) > max_chars:
        return (
            text[:max_chars].rstrip()
            + "\n\n[note] DOCX truncated: increase --max-chars or use start to read more."
        )
    return text
